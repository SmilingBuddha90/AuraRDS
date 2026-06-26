"""
Report Generator
Consolidates Analyze, Plan, and Operate pillar results into a formatted Word document.
"""

import os
import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)  # Finout blue
    elif level == 2:
        heading.runs[0].font.color.rgb = RGBColor(0x37, 0x4A, 0x5C)


def _add_pillar_section(doc: Document, pillar_name: str, pillar_title: str, content: str) -> None:
    _add_heading(doc, f"Pillar: {pillar_title}", level=1)
    doc.add_paragraph(f"Assessment: {pillar_name.upper()}", style="Intense Quote")
    doc.add_paragraph()

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("## "):
            _add_heading(doc, line[3:], level=2)
        elif line.startswith("### "):
            _add_heading(doc, line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("| "):
            doc.add_paragraph(line, style="No Spacing")
        else:
            doc.add_paragraph(line)


def generate_report(
    pillar_results: dict[str, str],
    query: str,
    reports_dir: str = "~/reports",
) -> str:
    """
    Generate a Word document report from pillar results.
    Returns the path to the generated file.
    """
    reports_path = Path(reports_dir).expanduser()
    reports_path.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"finout_rds_review_{timestamp}.docx"
    output_path = reports_path / filename

    doc = Document()

    # Cover page
    title = doc.add_heading("Finout RDS/Aurora Cost Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Query: {query}")
    doc.add_paragraph(f"Pillars assessed: {', '.join(pillar_results.keys()).upper()}")
    doc.add_paragraph()
    doc.add_page_break()

    # Executive Summary
    _add_heading(doc, "Executive Summary", level=1)
    doc.add_paragraph(
        "This report presents the results of an automated Finout-powered cost review "
        "of your RDS and Aurora database fleet across three pillars: Analyze (historical spend), "
        "Plan (forecasts and anomalies), and Operate (team ownership and showback)."
    )
    doc.add_paragraph()

    pillar_titles = {
        "analyze": "Analyze — Visibility into the Past",
        "plan":    "Plan — Confidence in the Future",
        "operate": "Operate — Ownership that Scales",
    }

    for pillar_key in ["analyze", "plan", "operate"]:
        if pillar_key in pillar_results:
            doc.add_page_break()
            _add_pillar_section(
                doc,
                pillar_key,
                pillar_titles[pillar_key],
                pillar_results[pillar_key],
            )

    # Documentation enrichment section
    if "docs_enrichment" in pillar_results:
        doc.add_page_break()
        _add_heading(doc, "Documentation References", level=1)
        doc.add_paragraph(
            "The following AWS and Finout documentation references support the "
            "recommendations made in this report."
        )
        doc.add_paragraph()
        for line in pillar_results["docs_enrichment"].split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith("## "):
                _add_heading(doc, line[3:], level=2)
            elif line.startswith("### "):
                _add_heading(doc, line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)

    # Footer
    doc.add_page_break()
    _add_heading(doc, "Report Notes", level=2)
    doc.add_paragraph(
        "This report was generated automatically using Finout MCP tools and Anthropic Claude "
        "via Amazon Bedrock. All cost figures are sourced directly from the Finout API. "
        "Recommendations should be validated before implementation."
    )
    doc.add_paragraph(f"Report saved: {output_path}")

    doc.save(str(output_path))
    logger.info(f"Report saved to: {output_path}")
    return str(output_path)
